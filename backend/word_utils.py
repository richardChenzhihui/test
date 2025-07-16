from docx import Document
import re
from typing import List, Dict, Optional

def parse_docx_structure(docx_path: str) -> Dict:
    doc = Document(docx_path)
    paragraphs = [
        {"index": i, "text": p.text}
        for i, p in enumerate(doc.paragraphs)
    ]
    tables = []
    for t_idx, table in enumerate(doc.tables):
        content = []
        for row in table.rows:
            content.append([cell.text for cell in row.cells])
        tables.append({
            "index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "content": content
        })
    return {"paragraphs": paragraphs, "tables": tables}

def edit_text(docx_path: str, locate_type: str, locate_value: str, op_type: str, content: Optional[str]) -> Dict:
    doc = Document(docx_path)
    para_idx = None
    if locate_type == 'index':
        para_idx = int(locate_value)
    elif locate_type == 'regex':
        for i, p in enumerate(doc.paragraphs):
            if re.search(locate_value, p.text):
                para_idx = i
                break
    elif locate_type == 'nl':
        # 预留大模型接口，暂用简单包含
        for i, p in enumerate(doc.paragraphs):
            if locate_value in p.text:
                para_idx = i
                break
    else:
        raise ValueError('未知定位方式')
    if para_idx is None or para_idx >= len(doc.paragraphs):
        return {"success": False, "message": "未找到目标段落"}
    para = doc.paragraphs[para_idx]
    if op_type == 'insert':
        para.text = para.text[:0] + (content or '') + para.text[0:]
    elif op_type == 'replace':
        para.text = content or ''
    elif op_type == 'delete':
        para.text = ''
    else:
        return {"success": False, "message": "未知操作类型"}
    save_path = docx_path.replace('.docx', '_edited.docx')
    doc.save(save_path)
    structure = parse_docx_structure(save_path)
    return {"success": True, "download_url": save_path, "structure": structure}