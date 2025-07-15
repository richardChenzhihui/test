from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from typing import List, Optional, Dict, Any
import os
import re
from pathlib import Path

from api.models import WordOperation, TextStyle, TableData, OperationType


class WordProcessor:
    """Word文档处理器"""
    
    def __init__(self, document_path: Optional[str] = None):
        """初始化文档处理器
        
        Args:
            document_path: 文档路径，如果为None则创建新文档
        """
        if document_path and os.path.exists(document_path):
            self.document = Document(document_path)
            self.document_path = document_path
        else:
            self.document = Document()
            self.document_path = None
    
    def save_document(self, output_path: str) -> bool:
        """保存文档
        
        Args:
            output_path: 输出路径
            
        Returns:
            是否成功保存
        """
        try:
            # 确保目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.document.save(output_path)
            self.document_path = output_path
            return True
        except Exception as e:
            print(f"保存文档失败: {e}")
            return False
    
    def add_text(self, content: str, position: str = "end", style: Optional[TextStyle] = None) -> bool:
        """添加文本
        
        Args:
            content: 文本内容
            position: 插入位置 ("end", "beginning", "after:段落索引")
            style: 文本样式
            
        Returns:
            是否成功
        """
        try:
            if position == "beginning":
                # 在开头插入
                paragraph = self.document.paragraphs[0]
                paragraph.insert_paragraph_before(content)
            elif position.startswith("after:"):
                # 在指定段落后插入
                index = int(position.split(":")[1])
                if index < len(self.document.paragraphs):
                    self.document.paragraphs[index].insert_paragraph_after(content)
                else:
                    self.document.add_paragraph(content)
            else:
                # 在末尾添加
                paragraph = self.document.add_paragraph(content)
                
            # 应用样式
            if style:
                self._apply_text_style(paragraph, style)
                
            return True
        except Exception as e:
            print(f"添加文本失败: {e}")
            return False
    
    def add_heading(self, content: str, level: int = 1, style: Optional[TextStyle] = None) -> bool:
        """添加标题
        
        Args:
            content: 标题内容
            level: 标题级别 (1-9)
            style: 文本样式
            
        Returns:
            是否成功
        """
        try:
            heading = self.document.add_heading(content, level)
            if style:
                self._apply_text_style(heading, style)
            return True
        except Exception as e:
            print(f"添加标题失败: {e}")
            return False
    
    def add_table(self, table_data: TableData, style: Optional[TextStyle] = None) -> bool:
        """添加表格
        
        Args:
            table_data: 表格数据
            style: 文本样式
            
        Returns:
            是否成功
        """
        try:
            table = self.document.add_table(rows=table_data.rows, cols=table_data.cols)
            table.style = 'Table Grid'
            
            # 添加表头
            if table_data.headers:
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data.headers):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = header
                        # 表头加粗
                        for paragraph in hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            # 添加数据
            if table_data.data:
                start_row = 1 if table_data.headers else 0
                for row_idx, row_data in enumerate(table_data.data):
                    if start_row + row_idx < len(table.rows):
                        row_cells = table.rows[start_row + row_idx].cells
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(row_cells):
                                row_cells[col_idx].text = str(cell_data)
            
            return True
        except Exception as e:
            print(f"添加表格失败: {e}")
            return False
    
    def modify_text(self, old_text: str, new_text: str, style: Optional[TextStyle] = None) -> bool:
        """修改文本
        
        Args:
            old_text: 原文本
            new_text: 新文本
            style: 文本样式
            
        Returns:
            是否成功
        """
        try:
            modified = False
            for paragraph in self.document.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    if style:
                        self._apply_text_style(paragraph, style)
                    modified = True
            
            return modified
        except Exception as e:
            print(f"修改文本失败: {e}")
            return False
    
    def delete_text(self, target_text: str) -> bool:
        """删除文本
        
        Args:
            target_text: 要删除的文本
            
        Returns:
            是否成功
        """
        try:
            paragraphs_to_remove = []
            for paragraph in self.document.paragraphs:
                if target_text in paragraph.text:
                    # 如果整个段落都是目标文本，则删除段落
                    if paragraph.text.strip() == target_text.strip():
                        paragraphs_to_remove.append(paragraph)
                    else:
                        # 否则只删除文本内容
                        paragraph.text = paragraph.text.replace(target_text, "")
            
            # 删除段落
            for paragraph in paragraphs_to_remove:
                self._delete_paragraph(paragraph)
            
            return True
        except Exception as e:
            print(f"删除文本失败: {e}")
            return False
    
    def add_list(self, items: List[str], ordered: bool = False, style: Optional[TextStyle] = None) -> bool:
        """添加列表
        
        Args:
            items: 列表项
            ordered: 是否有序列表
            style: 文本样式
            
        Returns:
            是否成功
        """
        try:
            for item in items:
                if ordered:
                    paragraph = self.document.add_paragraph(item, style='List Number')
                else:
                    paragraph = self.document.add_paragraph(item, style='List Bullet')
                
                if style:
                    self._apply_text_style(paragraph, style)
            
            return True
        except Exception as e:
            print(f"添加列表失败: {e}")
            return False
    
    def _apply_text_style(self, paragraph, style: TextStyle):
        """应用文本样式
        
        Args:
            paragraph: 段落对象
            style: 文本样式
        """
        try:
            for run in paragraph.runs:
                if style.font_name:
                    run.font.name = style.font_name
                if style.font_size:
                    run.font.size = Pt(style.font_size)
                if style.bold is not None:
                    run.font.bold = style.bold
                if style.italic is not None:
                    run.font.italic = style.italic
                if style.underline is not None:
                    run.font.underline = style.underline
                if style.color:
                    # 解析颜色（支持RGB格式）
                    if style.color.startswith('#'):
                        color_hex = style.color[1:]
                        if len(color_hex) == 6:
                            r = int(color_hex[0:2], 16)
                            g = int(color_hex[2:4], 16)
                            b = int(color_hex[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
        except Exception as e:
            print(f"应用样式失败: {e}")
    
    def _delete_paragraph(self, paragraph):
        """删除段落
        
        Args:
            paragraph: 要删除的段落
        """
        try:
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None
        except Exception as e:
            print(f"删除段落失败: {e}")
    
    def get_document_text(self) -> str:
        """获取文档全文
        
        Returns:
            文档文本内容
        """
        try:
            full_text = []
            for paragraph in self.document.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"获取文档文本失败: {e}")
            return ""
    
    def execute_operation(self, operation: WordOperation) -> bool:
        """执行单个操作
        
        Args:
            operation: 操作对象
            
        Returns:
            是否成功
        """
        try:
            if operation.operation_type == OperationType.ADD_TEXT:
                return self.add_text(operation.content, operation.position, operation.style)
            elif operation.operation_type == OperationType.DELETE_TEXT:
                return self.delete_text(operation.content)
            elif operation.operation_type == OperationType.MODIFY_TEXT:
                # 需要在metadata中提供old_text
                old_text = operation.metadata.get('old_text', '') if operation.metadata else ''
                return self.modify_text(old_text, operation.content, operation.style)
            elif operation.operation_type == OperationType.ADD_TABLE:
                return self.add_table(operation.table_data, operation.style)
            elif operation.operation_type == OperationType.ADD_HEADING:
                level = operation.metadata.get('level', 1) if operation.metadata else 1
                return self.add_heading(operation.content, level, operation.style)
            elif operation.operation_type == OperationType.ADD_LIST:
                items = operation.metadata.get('items', []) if operation.metadata else []
                ordered = operation.metadata.get('ordered', False) if operation.metadata else False
                return self.add_list(items, ordered, operation.style)
            else:
                print(f"不支持的操作类型: {operation.operation_type}")
                return False
        except Exception as e:
            print(f"执行操作失败: {e}")
            return False
    
    def execute_operations(self, operations: List[WordOperation]) -> List[bool]:
        """执行多个操作
        
        Args:
            operations: 操作列表
            
        Returns:
            每个操作的执行结果
        """
        results = []
        for operation in operations:
            result = self.execute_operation(operation)
            results.append(result)
        return results